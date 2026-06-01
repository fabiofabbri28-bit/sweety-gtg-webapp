"""
Generazione documenti in memoria:
- Lead Register (Excel)
- Activity Summary (Word)
- Bozza email Sales (txt)
"""
import io
import re
import calendar
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


MONTHS_EN = ["", "January", "February", "March", "April", "May", "June",
             "July", "August", "September", "October", "November", "December"]
MONTHS_ABBR = ["", "JAN", "FEB", "MAR", "APR", "MAY", "JUN",
               "JUL", "AUG", "SEP", "OCT", "NOV", "DEC"]
MONTHS_IT = ["", "Gennaio", "Febbraio", "Marzo", "Aprile", "Maggio", "Giugno",
             "Luglio", "Agosto", "Settembre", "Ottobre", "Novembre", "Dicembre"]

BLUE_DARK  = "1F4E79"
BLUE_MID   = "2E75B6"
BLUE_LIGHT = "DAE8F5"

_BLUE  = RGBColor(0x1F, 0x4E, 0x79)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ── Helpers generici ───────────────────────────────────────────────────────────

def _thin_border():
    s = Side(style="thin", color="CCCCCC")
    return Border(left=s, right=s, top=s, bottom=s)


def _fmt_euro(n: float) -> str:
    return f"€{n:,.0f}".replace(",", ".")


def _period_label(period: str) -> str:
    try:
        y, m = period.split("-")
        return f"{MONTHS_IT[int(m)]} {y}"
    except Exception:
        return period


def _period_label_en(period: str) -> str:
    try:
        y, m = period.split("-")
        return f"{MONTHS_EN[int(m)]} {y}"
    except Exception:
        return period


def _period_abbr(period: str) -> str:
    try:
        y, m = period.split("-")
        return f"{MONTHS_ABBR[int(m)]}-{y}"
    except Exception:
        return period


def _data_ym_to_en(data_ym: str) -> str:
    try:
        y, m = data_ym.split("-")
        return f"{MONTHS_EN[int(m)]} {y}"
    except Exception:
        return data_ym


def _period_full_range(period: str):
    """Returns ('1 May 2026', '31 May 2026') for a YYYY-MM string."""
    y, m = int(period[:4]), int(period[5:7])
    last = calendar.monthrange(y, m)[1]
    return f"1 {MONTHS_EN[m]} {y}", f"{last} {MONTHS_EN[m]} {y}"


def _origin_months_text(confirmed: list, period: str) -> str:
    """Returns the billing period month in English."""
    return _data_ym_to_en(period)


# ── Helpers Word ────────────────────────────────────────────────────────────────

def _word_cell_text(cell, text: str, bold=False, size=11,
                    color: RGBColor = None,
                    align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _remove_table_borders(table):
    """Rimuove tutti i bordi visibili da una tabella Word."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _add_para_bottom_border(para, hex_color=BLUE_DARK):
    """Aggiunge un bordo inferiore a un paragrafo."""
    pPr = para._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:pBdr")):
        pPr.remove(existing)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _fill_cell_lines(cell, lines: list, align=WD_ALIGN_PARAGRAPH.LEFT, size=11):
    """Riempie una cella con più righe di testo."""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        if i == 0:
            p.clear()
        p.alignment = align
        run = p.add_run(line)
        run.font.size = Pt(size)


def _add_section_heading(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = _BLUE
    run.font.size = Pt(12)
    return p


def _body_para(doc, text: str, space_after: int = 6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


# ── Lead Register (Excel) ──────────────────────────────────────────────────────

def generate_lead_register(confirmed: list, totals: dict,
                           period: str, invoice_sales: str) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Lead Register"

    hdr_fill  = PatternFill("solid", fgColor=BLUE_DARK)
    alt_fill  = PatternFill("solid", fgColor=BLUE_LIGHT)
    tot_fill  = PatternFill("solid", fgColor="FFF2CC")
    hdr_font  = Font(bold=True, color="FFFFFF", size=10)
    body_font = Font(size=10)
    bold_font = Font(bold=True, size=10)

    headers = ["Lead ID", "Origin period", "Customer", "City", "N. machines", "Billable"]
    ws.append(headers)
    for col in range(1, len(headers) + 1):
        c = ws.cell(1, col)
        c.fill = hdr_fill
        c.font = hdr_font
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = _thin_border()
    ws.row_dimensions[1].height = 22

    prefix = _period_abbr(period)
    seq = 0
    for cl in confirmed:
        for row_data in cl.get("rows", []):
            seq += 1
            lead_id = f"{prefix}-{seq:03d}"
            origin  = _data_ym_to_en(row_data.get("data_ym", period))
            address = row_data.get("indirizzo", "")   # full address in City column
            ws.append([lead_id, origin, cl["nome"], address, 1, "YES"])
            r = ws.max_row
            fill = alt_fill if seq % 2 == 0 else PatternFill()
            for col in range(1, 7):
                c = ws.cell(r, col)
                c.fill = fill
                c.font = body_font
                c.border = _thin_border()
                c.alignment = Alignment(
                    horizontal="left" if col in (3, 4) else "center",
                    vertical="center",
                )
            ws.row_dimensions[r].height = 18

    # Riga totale
    total_machines = sum(len(c.get("rows", [])) for c in confirmed)
    ws.append(["", "", "", "Total", total_machines, ""])
    r = ws.max_row
    for col in range(1, 7):
        c = ws.cell(r, col)
        c.fill = tot_fill
        c.font = bold_font
        c.border = _thin_border()
        c.alignment = Alignment(horizontal="center", vertical="center")

    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 14
    ws.column_dimensions["C"].width = 38
    ws.column_dimensions["D"].width = 55
    ws.column_dimensions["E"].width = 14
    ws.column_dimensions["F"].width = 12

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── Activity Summary (Word) ────────────────────────────────────────────────────

def generate_activity_summary(totals: dict, period: str,
                               sweety: dict, gtg: dict,
                               confirmed: list = None) -> bytes:
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(3.0)

    first_date, last_date = _period_full_range(period)
    period_range = f"{first_date} – {last_date}"
    origins_text = _origin_months_text(confirmed, period)
    sweety_name = sweety.get("name", "Sweety Pact S.r.l.")
    gtg_name    = gtg.get("name", "Good to Great S.r.l.")

    # ── Titolo ──────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    tr = title_p.add_run("ACTIVITY SUMMARY REPORT")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = _BLUE
    _add_para_bottom_border(title_p)

    doc.add_paragraph()

    # ── Period line ─────────────────────────────────────────────────────────────
    _body_para(doc, f"Period: {period_range}")

    doc.add_paragraph()

    # ── Service Provider / Client (2-col borderless table) ─────────────────────
    pc = doc.add_table(rows=1, cols=2)
    _remove_table_borders(pc)

    _fill_cell_lines(pc.cell(0, 0), [
        "Service Provider:",
        sweety_name,
        sweety.get("address_line1", ""),
        sweety.get("address_line2", ""),
        f"Cod Fiscal {sweety.get('cod_fiscal', '')}",
    ])

    _fill_cell_lines(pc.cell(0, 1), [
        "Client:",
        gtg_name,
        gtg.get("address", "România, comuna Arad, str. Poetului, nr.1/c, hala 21"),
        f"Cod Fiscal {gtg.get('cod_fiscal', '')}",
    ], align=WD_ALIGN_PARAGRAPH.RIGHT)

    for row in pc.rows:
        row.cells[0].width = Cm(8.5)
        row.cells[1].width = Cm(8.5)

    doc.add_paragraph()

    # ── Reference period + Report date ──────────────────────────────────────────
    _body_para(doc, "Reference period:", space_after=0)
    _body_para(doc, period_range)

    doc.add_paragraph()

    _body_para(doc, "Report date:", space_after=0)
    _body_para(doc, last_date)

    doc.add_paragraph()

    # ── Section 1: Purpose ──────────────────────────────────────────────────────
    _add_section_heading(doc, "1. Purpose of the report")
    _body_para(doc,
        f"This report summarizes the activities performed by {sweety_name} in favor of "
        f"{gtg_name} during the above-mentioned reference period, within the scope of the "
        f"client acquisition and business development support agreement in force between the "
        f"parties. The document is issued as supporting documentation for invoicing and "
        f"accounting purposes.")

    doc.add_paragraph()

    # ── Section 2: Description ──────────────────────────────────────────────────
    _add_section_heading(doc, "2. Description of activities performed")
    _body_para(doc,
        f"During the reference period, {sweety_name} carried out client acquisition and "
        f"business development support activities, including in particular:", space_after=4)

    for item in [
        "identification and qualification of prospective customers;",
        "follow-up activities and commercial support;",
        "coordination of customer onboarding and activation processes;",
        "operational support related to customer acquisition initiatives.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"- {item}")
        run.font.size = Pt(11)

    doc.add_paragraph()
    _body_para(doc,
        "All activities were performed remotely from the Republic of Moldova, in accordance "
        "with the contractual provisions.")

    doc.add_paragraph()

    # ── Section 3: Quantitative summary ─────────────────────────────────────────
    _add_section_heading(doc, "3. Quantitative summary of activities")

    tbl = doc.add_table(rows=4, cols=4)
    _remove_table_borders(tbl)

    # Header row
    for i, h in enumerate(["Service category", "Quantity", "Unit value", "Amount"]):
        _word_cell_text(tbl.cell(0, i), h, bold=True, size=11)

    # Fixed fee row
    _word_cell_text(tbl.cell(1, 0), "Fixed monthly fee – client acquisition support")
    _word_cell_text(tbl.cell(1, 1), "1")
    _word_cell_text(tbl.cell(1, 2), _fmt_euro(totals["fixed_fee"]))
    _word_cell_text(tbl.cell(1, 3), _fmt_euro(totals["fixed_fee"]))

    # Operational activations row
    _word_cell_text(tbl.cell(2, 0), "Operational activations")
    _word_cell_text(tbl.cell(2, 1), str(totals["macchine"]))
    _word_cell_text(tbl.cell(2, 2), _fmt_euro(500))
    _word_cell_text(tbl.cell(2, 3), _fmt_euro(totals["variabile"]))

    # Total row
    _word_cell_text(tbl.cell(3, 0), "Total services for the reference period", bold=True)
    _word_cell_text(tbl.cell(3, 1), "")
    _word_cell_text(tbl.cell(3, 2), "")
    _word_cell_text(tbl.cell(3, 3), _fmt_euro(totals["totale_sales"]), bold=True)

    for row in tbl.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(2)
        row.cells[2].width = Cm(2.5)
        row.cells[3].width = Cm(2.5)

    doc.add_paragraph()

    # ── Section 4: Measurement and cut-off ──────────────────────────────────────
    _add_section_heading(doc, "4. Measurement and cut-off criteria")
    _body_para(doc,
        "Operational activations represent completed activity units relevant for remuneration "
        "purposes, in accordance with the contractually defined criteria.")

    _body_para(doc,
        f"The reporting period includes activities originated during {origins_text} "
        "and accounted for in accordance with the standard operational reconciliation "
        "procedures adopted between the parties.")

    _body_para(doc,
        "All activities included in this report have been recognized in accordance with the "
        "contractual criteria applicable to the reporting period.")

    # ── Footer ──────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    for i, line in enumerate([
        sweety_name,
        sweety.get("address_line1", ""),
        sweety.get("address_line2", ""),
        f"Cod Fiscal {sweety.get('cod_fiscal', '')}",
        sweety.get("email", "info@sweetypact.com"),
        sweety.get("phone", "+37360045404"),
    ]):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.size = Pt(11)
        if i == 0:
            run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Bozza email Sales (txt) ────────────────────────────────────────────────────

def generate_email_draft_sales(totals: dict, period: str,
                                invoice_sales: str,
                                sweety: dict, gtg: dict) -> bytes:
    period_en = _period_label_en(period)
    body = (
        f"Good morning,\n\n"
        f"as a follow-up to the sales support services provided, please find attached "
        f"the Activity Summary Report and the related invoice for the month of {period_en}, "
        f"prepared in line with the contractual terms in force.\n\n"
        f"The report is provided as supporting documentation for invoicing and reconciliation "
        f"purposes.\n\n"
        f"The reporting period includes activities accounted for in accordance with the "
        f"contractual criteria and the operational reconciliation procedures adopted between "
        f"the parties.\n\n"
        f"Should you need any further clarification, we remain at your disposal.\n\n"
        f"Kind regards,\n\n"
        f"{sweety.get('name', 'Sweety Pact S.r.l.')}\n"
        f"{sweety.get('email', 'info@sweetypact.com')}\n"
        f"{sweety.get('phone', '+37360045404')}\n"
    )
    return body.encode("utf-8")


# ── Genera tutti i documenti ───────────────────────────────────────────────────

def generate_all(confirmed: list, totals: dict, period: str,
                 invoice_sales: str, invoice_admin: str,
                 sweety: dict, gtg: dict, admin_contact: dict) -> dict:
    return {
        f"lead_register_{period}.xlsx": generate_lead_register(
            confirmed, totals, period, invoice_sales),
        f"activity_summary_{period}.docx": generate_activity_summary(
            totals, period, sweety, gtg, confirmed),
        f"email_sales_{period}.txt": generate_email_draft_sales(
            totals, period, invoice_sales, sweety, gtg),
    }
